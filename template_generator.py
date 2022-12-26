from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

"""
    How to run:
        python template_generator.py
"""

DOCX_EXTENSION: str = ".docx"
DATE_LABEL: str = "Data:"
MODE_LABEL: str = "Tryb:"
SUMMARY_LABEL: str = "Podsumowanie:"
DESCRIPTION_LABEL: str = "Opis:"


def main() -> None:
    document = Document()

    date_paragraph = document.add_paragraph()
    date_paragraph.add_run(f"{DATE_LABEL}").bold = True
    date_paragraph.add_run(" ")
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    mode_paragraph = document.add_paragraph()
    mode_paragraph.add_run(f"{MODE_LABEL} ").bold = True
    mode_paragraph.add_run(" ")
    mode_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    document.add_heading("Sprawozdanie", level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    description_paragraph = document.add_paragraph()
    description_paragraph.add_run(f"{DESCRIPTION_LABEL}").bold = True
    description_paragraph.add_run(" ")

    summary_paragraph = document.add_paragraph()
    summary_paragraph.add_run(f"{SUMMARY_LABEL}").bold = True
    summary_paragraph.add_run(" ")

    document.save(f"kollaps_report_XX-XX-XXXX{DOCX_EXTENSION}")


if __name__ == "__main__":
    main()
