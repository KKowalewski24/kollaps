import glob
from argparse import ArgumentParser, Namespace
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_BREAK

"""
    How to run:
        
"""

DOCX_EXTENSION: str = ".docx"
OUTPUT_FILENAME: str = "summary"
DATE_LABEL: str = "Data:"
MODE_LABEL: str = "Tryb:"
SUMMARY_LABEL: str = "Podsumowanie:"


def main() -> None:
    args = prepare_args()
    path: str = args.path
    filepaths: List[str] = glob.glob(f"{path}*{DOCX_EXTENSION}")

    output_document = Document()
    for filepath in filepaths:
        input_document: Document = Document(filepath)
        date, mode, summary = extract_information(input_document)

        output_document.add_paragraph(f"{DATE_LABEL} {date}")
        output_document.add_paragraph(f"{MODE_LABEL} {mode}")
        output_document.add_paragraph(f"{SUMMARY_LABEL} {summary}")
        output_document.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    output_document.save(f"{OUTPUT_FILENAME}{DOCX_EXTENSION}")


def extract_information(input_document: Document) -> Tuple[str, str, str]:
    date = ""
    mode = ""
    summary = ""

    for paragraph in input_document.paragraphs:
        if DATE_LABEL in paragraph.text:
            date = str(paragraph.text).replace(DATE_LABEL, "")
        if MODE_LABEL in paragraph.text:
            mode = str(paragraph.text).replace(MODE_LABEL, "")
        if SUMMARY_LABEL in paragraph.text:
            summary = str(paragraph.text).replace(SUMMARY_LABEL, "")

    return date, mode, summary


def prepare_args() -> Namespace:
    arg_parser = ArgumentParser()
    arg_parser.add_argument(
        "-p", "--path", type=str, required=True,
        help="Path do directory where reports are placed"
    )
    return arg_parser.parse_args()


if __name__ == "__main__":
    main()
