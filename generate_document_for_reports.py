from datetime import datetime

from docx import Document
from docx.enum.text import WD_BREAK

DOCX_EXTENSION: str = ".docx"
DATE_LABEL: str = "Data:"
MODE_LABEL: str = "Tryb:"
SUMMARY_LABEL: str = "Podsumowanie:"


def main() -> None:
    dates = [
        "12.12.2022",
        "05.05.2021",
        "24.12.2021",
        "26.12.2022",
    ]

    dates = [datetime.strptime(date, "%d.%m.%Y") for date in dates]
    dates.sort()
    dates = [f"{date.day}.{date.strftime('%m')}.{date.year}" for date in dates]

    document = Document()
    for date in dates:
        date_paragraph = document.add_paragraph()
        date_paragraph.add_run(f"{DATE_LABEL}").bold = True
        date_paragraph.add_run(f" {date}r")

        mode_paragraph = document.add_paragraph()
        mode_paragraph.add_run(f"{MODE_LABEL}").bold = True
        mode_paragraph.add_run(" ")

        summary_paragraph = document.add_paragraph()
        summary_paragraph.add_run(f"{SUMMARY_LABEL}").bold = True
        summary_paragraph.add_run(" ")

        document.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    document.save(f"podsumowanie{DOCX_EXTENSION}")


if __name__ == "__main__":
    main()
